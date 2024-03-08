import re
from re import Match
from typing import Optional, Final

import docx
from docx.document import Document
from docx.shared import Pt


class Converter:
    def __init__(self, doc_path: str, result_path: str) -> None:
        self._original: Document = docx.Document(doc_path)
        self._result_path: str = result_path
        self._REGEX: Final[str] = r'\$\{[^${}]*\}'  # ${var}

    @staticmethod
    def _get_str_match(accordance: Match) -> str:
        return accordance.string[accordance.span()[0]:accordance.span()[1]]

    @staticmethod
    def _get_name_match(str_accordance: str) -> str:
        return str_accordance[2:-1]

    def replace_vars(self, dict_vars: dict[str, str]) -> None:
        for paragraph in self._original.paragraphs:
            size = paragraph.runs[0].font.size.pt
            for match in re.finditer(self._REGEX, paragraph.text):
                str_match: str = self._get_str_match(match)
                name: str = self._get_name_match(str_match)
                if name in dict_vars:
                    paragraph.text = paragraph.text.replace(str_match, dict_vars[name])
            paragraph.runs[0].font.size = Pt(size)
        self._original.save(self._result_path)

    def get_vars(self) -> list[str]:
        result: list[str] = []
        for paragraph in self._original.paragraphs:
            for var in re.findall(self._REGEX, paragraph.text):
                str_var = self._get_name_match(var)
                result.append(str_var)
        return result
